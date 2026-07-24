import docx
from pathlib import Path
from typing import Dict, Any
from app.core.logging import logger
from app.documents.base_processor import BaseDocumentProcessor
from app.documents.exceptions import DocumentCorruptedException, DocumentProcessingException


class DOCXProcessor(BaseDocumentProcessor):
    """Memory-efficient Microsoft Word (.docx) document parser utilizing python-docx."""

    def extract_text(self, file_path: Path) -> Dict[str, Any]:
        logger.info(f"Initiating DOCX extraction pipeline for: {file_path.name}")
        
        try:
            doc = docx.Document(file_path)
        except Exception as exc:
            logger.error(f"Failed to open DOCX document: {exc}", exc_info=True)
            raise DocumentCorruptedException(message="Word document (.docx) appears corrupted or unreadable")

        paragraphs_text = []
        char_count = 0
        word_count = 0

        try:
            # 1. Extract Paragraph Text
            for para in doc.paragraphs:
                text = para.text or ""
                if text.strip():
                    paragraphs_text.append(text)
                    char_count += len(text)
                    word_count += len(text.split())

            # 2. Extract Table cell text for complete extraction
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text or ""
                        if text.strip():
                            paragraphs_text.append(text)
                            char_count += len(text)
                            word_count += len(text.split())

            full_text = "\n\n".join(paragraphs_text)

            logger.info(
                f"Successfully parsed Word DOCX: {file_path.name}",
                extra={
                    "extra": {
                        "character_count": char_count,
                        "word_count": word_count
                    }
                }
            )

            return {
                "text": full_text,
                "pages": None,  # Pages not natively stored in DOCX XML structure
                "character_count": char_count,
                "word_count": word_count,
            }

        except Exception as exc:
            logger.error(f"Error extracting text from DOCX elements: {exc}", exc_info=True)
            raise DocumentProcessingException(message="Internal DOCX text extraction failed")
